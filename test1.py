from docx import Document
from os.path import join

docx = Document("C:\\Users\\LENOVO\\OneDrive\\Desktop\\Test.docx")
saving_path = "C:\\Users\\LENOVO\\OneDrive\\Desktop"
# saving_path= "C:\\Users\\LENOVO\\OneDrive\\Desktop"
nsmap = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
}
# gonna append ques_no and image_name
ques_image={}
run = docx.paragraphs[1].runs[-1]
image_object_blip = run._element.findall('.//a:blip', namespaces=nsmap)
if image_object_blip:
    print("This is Image. Writing image to desktop")
    for obj_blip in image_object_blip:
        rId = obj_blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rId:
            image_part = docx.part.rels[rId].target_part
            image_binary_data = image_part.blob
            image_format = image_part.content_type.split("/")[-1]
            image_path = join(saving_path,f"adminname-quesno.{image_format}")
            with open(image_path,"wb") as f : # type: ignore
                f.write(image_binary_data)




print(len(docx.paragraphs))

# print(docx.paragraphs[0].runs)
