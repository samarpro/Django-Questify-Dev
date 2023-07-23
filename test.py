from docx import Document
from random import sample
from lxml import etree
# def WordFileHandler(FilePath,Quesno):
#     ANS_LIST = []
#     PARSE_DICT ={}
#     doc = Document(FilePath)
#     paraObj = doc.paragraphs
#     noPara = len(paraObj)
#     actNoPara = (noPara//5)*5
#     randomList = sample(range(0,actNoPara,5),Quesno)
#     for num in randomList:
#         paraText = paraObj[num].text
#         optionRandomList=sample(range(1,5),4)
#         OPTION_LIST = []
#         for opNum in optionRandomList:
#             optionObject = paraObj[num+opNum]
#             if optionObject.runs[0].bold == True:
#                 ANS_LIST.append(len(OPTION_LIST))
#             OPTION_LIST.append(optionObject.text)
        
#         PARSE_DICT[paraText] = OPTION_LIST
#     PARSE_DICT["Answers"] = ANS_LIST

#     return PARSE_DICT
    
    
# WordFileHandler('E:/Django-Questify-Dev/media/WordFiles/MCQ.docx',40)

docx = Document("C:\\Users\\LENOVO\\OneDrive\\Desktop\\Test.docx")
saving_path= "C:\\Users\\LENOVO\\OneDrive\\Desktop"


# print(docx.paragraphs[0].text)
print(len(docx.paragraphs))
# prints no of images in file.
print(len(docx.inline_shapes))



import os
from docx import Document

def extract_and_remove_images(docx_file, save_directory,paragraph):
    if not os.path.exists(save_directory):
        os.makedirs(save_directory)

    doc = Document(docx_file)


    # List to store the filenames of the saved images
    saved_image_filenames = []

    for rel in doc.part.rels:
        part = doc.part.rels[rel].target_part
        if "image" in part.content_type:
            # Extract the image binary data
            image_binary = part.blob

            # Get the filename from the part name
            image_filename = part.partname.split("/")[-1]

            # Save the image to the desired location
            image_path = os.path.join(save_directory, image_filename)
            with open(image_path, "wb") as f:
                f.write(image_binary)

            # Add the image part and saved image filename to the respective lists
            saved_image_filenames.append(image_filename)
            print(saved_image_filenames)
            # delete that specific image
            # doc.part.rels.pop(rel)
    # Save the modified document without the images
    doc.save(docx_file)

    return saved_image_filenames


# checks wheather the paragragh contains image or not
def identify_content(paragraph):
    paragraph_element = etree.fromstring(paragraph._element.xml.encode()) #type:ignore
    # used to uniquely identify elements and attributes in XML
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    }
    image_ele=paragraph_element.xpath('.//w:drawing',namespaces=namespaces)
    if image_ele:
        saved_images = extract_and_remove_images("C:\\Users\\LENOVO\\OneDrive\\Desktop\\Test.docx",saving_path,paragraph)
        print('IMAGE')
        print("Extracted and removed images:", saved_images)
    else:
        print(paragraph.text)

identify_content(docx.paragraphs[1])