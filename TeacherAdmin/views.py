from django.shortcuts import render,redirect
from django.http import JsonResponse
from .forms import AdminInfoForms
from .models import AdminInfo
from validator.models import CustomUser
from pathlib import Path
from django.conf import settings
from StPage.models import StInfoModels
from docx import Document
from os import remove
from datetime import datetime
# Create your views here.
def dashboard(req):
    user=req.user
    # this is dict is used to store which object is active of different grade
    context_of_object ={}
    if not user.is_anonymous:
        AdminInfo_objects = AdminInfo.objects.filter(user=user)
        St_object_Pass = len(StInfoModels.objects.filter(Pass=True,User=user))
        St_object_Total = len(StInfoModels.objects.filter(User=user))
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M')
        for grade in AdminInfo_objects.values('FOR_CLASS').distinct().order_by('FOR_CLASS'):
            # The Admin Info object which is currently active.
            # since last one is always taken as active 
            str_grade =str(grade['FOR_CLASS'])
            if str_grade in req.session:
                id_adminInfo = req.session[str_grade]
                active_admin_object = AdminInfo_objects.get(id=id_adminInfo) 
            else:
                active_admin_object = AdminInfo_objects.filter(FOR_CLASS = grade['FOR_CLASS']).last() 
                
            context_of_object[f"Grade{grade['FOR_CLASS']}"]=active_admin_object
        context={
            "AdminInfo_objects": AdminInfo_objects,
            'current_time': current_time,
            'Total_student':St_object_Total,
            'Pass_student':St_object_Pass,
            'active_list':context_of_object,
        }
        return render(req,"TeacherAdmin/dashboard.html",context)
    else:
        return redirect("validator/")

def uploadInfo(req):
    err=msg=""
    if req.method=="POST":
        form =  AdminInfoForms(req.POST,req.FILES)
        if form.is_valid():
            username = req.user
            InsName=form.cleaned_data['INSNAME']
            fmarks= form.cleaned_data['FULLMARKS']
            pmarks= form.cleaned_data['PASSMARKS']
            file=form.cleaned_data['FILENAME']
            addtext= form.cleaned_data['ADDTEXT']
            for_class=form.cleaned_data['FOR_CLASS']
            start_time =form.cleaned_data['Start_time']
            end_time =form.cleaned_data['End_time']
            # ActUser = User.objects.get(username=username)

            # opening a new document
            Docx_file = Document(file)
            if fmarks<=len(Docx_file.paragraphs)//5:

                Obj = AdminInfo(user=username,INSNAME=InsName,FULLMARKS=fmarks,PASSMARKS=pmarks,FILENAME=file,ADDTEXT=addtext,FOR_CLASS=for_class,Start_time=start_time,End_time=end_time)
                Obj.save() # type: ignore
                msg = "Data saved and deployed to Database."
                err= False
                # namespace for finding images

                CustomNsmap = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                }
                dict_of_images_and_ques = {}
                "Checking each paragraph if contains images at last or not."
                for index,paragraph in enumerate(Docx_file.paragraphs):
                    image_object_blip=paragraph.runs[-1]._element.findall('.//a:blip', namespaces=CustomNsmap)
                    if image_object_blip:
                        static_images_root = Path(settings.MEDIA_ROOT) / "Admin_Images"
                        for obj_blip in image_object_blip:
                            rId = obj_blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rId:
                                image_part = Docx_file.part.rels[rId].target_part
                                image_binary_data = image_part.blob
                                image_format = image_part.content_type.split("/")[-1]
                                # image_path = join(saving_path,f"adminname-quesno.{image_format}")
                                image_name = f"{InsName}-{Docx_file.core_properties.title}-{index}.{image_format}"
                                # updating dict_of_images_and_ques. This is a dict which stores question number and its corresponding image name
                                dict_of_images_and_ques[index]=image_name
                                image_path = Path(static_images_root)/image_name
                                with open(image_path,"wb") as f : # type: ignore
                                    f.write(image_binary_data)
                Obj.IMAGES_DICT=dict_of_images_and_ques
                Obj.save()
                msg = "File Successfully Deployed."
                # wordFile = Document(BASE_DIR / f"media\WordFiles\{file.name}") #type:ignore
                # return HttpResponse(len(wordFile.paragraphs))
            else:
                err="No of Question required is less than Found in WordFile"        
        else:
            err = "Form Invalid. Please try again."
            # gets redirected to upload Page with errors.
        return render(req,"TeacherAdmin/upload.html",context={'err':err,'form':form,'msg':msg})
    else:
        form = AdminInfoForms()
        context = {
            'form':form,
            'err':err,
            'msg':msg
        }
        return render(req,"TeacherAdmin/upload.html",context)


def show_result(req):
    # this user refers to the Admin Name
    user=req.user
    if not user.is_anonymous:
        AdminUser = AdminInfo.objects.filter(user=user).last()
        if AdminUser is not None:
            # all the student of specific user.
            List_St = StInfoModels.objects.filter(User=user)
            if List_St.exists():
                context={
                    "ListSt":List_St,
                    "Student_exists_msg": None
                }
            else:
                context={
                    "Student_exists_msg": "No Student exists yet."
                }
        else:
            # if Admin User doesn't exists
            context={
                "Student_exists_msg":"Admin doesn't exists."
            }
        return render(req,"TeacherAdmin/result.html",context)
    else:
        return redirect("validator/")
    

def delete_row(req,id_of_row):
    model_row = AdminInfo.objects.get(id=int(id_of_row))
    path_wordFile = Path(settings.MEDIA_ROOT) / model_row.FILENAME.name
    print(path_wordFile)
    """
    Institution name--{key_of_dict}
    """
    # model row -> object of admin info

    try:
        for val in model_row.IMAGES_DICT.values() :
            remove(Path(settings.MEDIA_ROOT)/"Admin_Images"/val)
        remove(path_wordFile)
    except Exception as e:
        print(e)
        return JsonResponse({'msg':False})

    model_row.delete()
    return JsonResponse({'msg':True})


def edit_object_info(req,id_object):
    err=msg=""
    try:
        Obj = AdminInfo.objects.get(id=id_object)
    except:
        return render(req,"TeacherAdmin/error_page.html") 
    if req.method=="POST":
        form =  AdminInfoForms(req.POST,req.FILES)
        form.fields['FILENAME'].required = False # to provide user a option of uploading and not uploading
        if form.is_valid():
            username = req.user
            InsName=form.cleaned_data['INSNAME']
            fmarks= form.cleaned_data['FULLMARKS']
            pmarks= form.cleaned_data['PASSMARKS']
            file=form.cleaned_data['FILENAME']
            addtext= form.cleaned_data['ADDTEXT']
            for_class=form.cleaned_data['FOR_CLASS']
            start_time =form.cleaned_data['Start_time']
            end_time =form.cleaned_data['End_time']
            # ActUser = User.objects.get(username=username)
            # opening a new document
            #  since Uploading a new file is optional
            if file is not None:            
                Docx_file = Document(file)
                if fmarks<=len(Docx_file.paragraphs)//5:
                    # Highlights upon the object with given id. 
                    Obj.FOR_CLASS = for_class
                    Obj.INSNAME = InsName
                    Obj.FULLMARKS = fmarks
                    Obj.PASSMARKS = pmarks
                    Obj.FILENAME = file
                    Obj.ADDTEXT = addtext
                    Obj.Start_time = start_time
                    Obj.End_time = end_time                
                    Obj.save() # type: ignore
                    err= False

                    # namespace for finding images
                    CustomNsmap = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                    }
                    dict_of_images_and_ques = {}
                    "Checking each paragraph if contains images at last or not."
                    for index,paragraph in enumerate(Docx_file.paragraphs):
                        image_object_blip=paragraph.runs[-1]._element.findall('.//a:blip', namespaces=CustomNsmap)
                        if image_object_blip:
                            static_images_root = Path(settings.MEDIA_ROOT) / "Admin_Images"
                            for obj_blip in image_object_blip:
                                rId = obj_blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rId:
                                    image_part = Docx_file.part.rels[rId].target_part
                                    image_binary_data = image_part.blob
                                    image_format = image_part.content_type.split("/")[-1]
                                    # image_path = join(saving_path,f"adminname-quesno.{image_format}")
                                    image_name = f"{InsName}-{Docx_file.core_properties.title}-{index}.{image_format}"
                                    # updating dict_of_images_and_ques. This is a dict which stores question number and its corresponding image name
                                    dict_of_images_and_ques[index]=image_name
                                    image_path = Path(static_images_root)/image_name
                                    with open(image_path,"wb") as f : # type: ignore
                                        f.write(image_binary_data)
                    Obj.IMAGES_DICT=dict_of_images_and_ques
                    Obj.save()
                    # messages
                    msg = "Data saved and deployed to Database."

                    # wordFile = Document(BASE_DIR / f"media\WordFiles\{file.name}") #type:ignore


                    # return HttpResponse(len(wordFile.paragraphs))
                else:
                    # if noo fo question is invalid
                    err="No of Question required is less than Found in WordFile"   
            else:
                # Even though file is not uploaded you still need to update the database for other fields.
                Obj.FOR_CLASS = for_class
                Obj.INSNAME = InsName
                Obj.FULLMARKS = fmarks
                Obj.PASSMARKS = pmarks
                Obj.ADDTEXT = addtext
                Obj.Start_time = start_time
                Obj.End_time = end_time                
                Obj.save() # type: ignore
                msg = "Data saved and deployed to Database."
                err= False

        else:
            err = "Form Invalid. Please try again."
            # gets redirected to upload Page with errors.
        return render(req,"TeacherAdmin/edit.html",context={'err':err,'form':form,'msg':msg})
    else:
        form = AdminInfoForms()
        # setting up initial value for forms
        # since we need to show value at beginning when req.method is not POST 
        form.fields['FOR_CLASS'].initial = Obj.FOR_CLASS
        form.fields['INSNAME'].initial = Obj.INSNAME
        form.fields['FULLMARKS'].initial = Obj.FULLMARKS
        form.fields['PASSMARKS'].initial = Obj.PASSMARKS
        form.fields['FILENAME'].initial = Obj.FILENAME
        form.fields['ADDTEXT'].initial = Obj.ADDTEXT
        form.fields['Start_time'].initial = Obj.Start_time
        form.fields['End_time'].initial = Obj.End_time
        context = {
            'form':form,
            'err':err,
            'msg':msg,
            'file_content':Obj.FILENAME
        }
        return render(req,"TeacherAdmin/edit.html",context)

    return render(req,'TeacherAdmin/edit.html')


def change_active_portal(req,id_object,class_id):

    """Takes ID of object and class of that object in DB."""
    user_object = CustomUser.objects.get(username=str(req.user))
    user_object.active_portals[class_id] = id_object  #type:ignore
    user_object.save()
    req.session[str(class_id)] = id_object
    return JsonResponse({'msg':id_object,'class':class_id})